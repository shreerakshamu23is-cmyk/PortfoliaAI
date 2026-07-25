import io
import re
from pypdf import PdfReader

import pdfplumber
import docx

try:
    import pytesseract
    from PIL import Image
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False

class ResumeParserService:
    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        """Extract text content from PDF bytes using pdfplumber with Tesseract OCR fallback."""
        text = ""
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    extracted = page.extract_text(layout=True)
                    if not extracted:
                        extracted = page.extract_text()
                    
                    # OCR Fallback for scanned image pages
                    if (not extracted or len(extracted.strip()) < 30) and TESSERACT_AVAILABLE:
                        try:
                            p_img = page.to_image(resolution=150).original
                            ocr_text = pytesseract.image_to_string(p_img)
                            if ocr_text:
                                extracted = ocr_text
                        except Exception as ocr_err:
                            print(f"[Parser] Page OCR fallback error: {ocr_err}")

                    if extracted:
                        text += extracted + "\n"
        except Exception as e:
            print(f"[Parser] pdfplumber failed: {e}. Falling back to pypdf...")
            try:
                reader = PdfReader(io.BytesIO(file_bytes))
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
            except Exception as ex:
                raise ValueError(f"Could not parse PDF file: {ex}")
                
        return text.strip()

    @staticmethod
    def extract_text_from_image(file_bytes: bytes) -> str:
        """Extract text content from image bytes (PNG, JPG, JPEG) using pytesseract OCR."""
        if not TESSERACT_AVAILABLE:
            raise ValueError("Tesseract OCR is not installed for processing image resumes.")
        try:
            image = Image.open(io.BytesIO(file_bytes))
            text = pytesseract.image_to_string(image)
            return text.strip()
        except Exception as e:
            raise ValueError(f"Could not parse image resume: {e}")

    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        """Extract text content from DOCX bytes using python-docx."""
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))

            return "\n".join(full_text)
        except Exception as e:
            raise ValueError(f"Could not parse DOCX file: {e}")

    @staticmethod
    def sanitize_pdf_raw_text(text: str) -> str:
        """Sanitizes raw PDF stream markers, font width arrays, hex strings, and decodes /Title annotations cleanly."""
        if not text:
            return ""

        # 1. Strip PDF Font Width Arrays (e.g. 667 722 722 1000 or [667 722 722])
        text = re.sub(r'\b\d{3,4}(?:\s+\d{3,4}){2,}\b', '', text)
        text = re.sub(r'\[\s*\d+(?:\s+\d+)*\s*\]', '', text)

        # 2. Decode UTF-16BE Hex titles (e.g. /Title <FEFFF0D80020...>)
        def decode_hex_title(m):
            hex_str = m.group(1)
            try:
                if hex_str.startswith('FEFF') or hex_str.startswith('feff'):
                    hex_str = hex_str[4:]
                decoded = bytes.fromhex(hex_str).decode('utf-16be', errors='ignore')
                return f'/Title ({decoded})'
            except Exception:
                return m.group(0)

        text = re.sub(r'/Title\s*<([0-9A-Fa-f]+)>', decode_hex_title, text)

        # 3. Extract /Title (...) annotations if present
        extracted_titles = []
        for m in re.finditer(r'/Title\s*\((.*?)\)', text, re.DOTALL):
            clean_t = m.group(1).replace(r'\(', '(').replace(r'\)', ')').strip()
            if clean_t and not any(junk in clean_t for junk in ['%PDF-', '/Dest', '/Parent', '/Prev', '/Next']):
                extracted_titles.append(clean_t)

        # 4. Filter out PDF binary stream junk & font specs
        pdf_junk_pattern = re.compile(
            r'(%PDF-|obj\b|endobj\b|/Dest\b|/Parent\b|/Prev\b|/Next\b|/XYZ\b|/Nums\b|/Footnote\b|/Endnote\b|'
            r'/Textbox\b|/Header\b|/Footer\b|/InlineShape\b|/Annotation\b|/Artifact\b|/Workbook\b|/Worksheet\b|'
            r'/Sect\b|/Document\b|/Part\b|<<|>>|\b\d+\s+\d+\s+R\b|Arial,Bold|TimesNewRoman|Helvetica|/Font\b|/Type\b)',
            re.IGNORECASE
        )

        clean_lines = []
        for line in text.split('\n'):
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith('%') or stripped.startswith('\ufffd') or stripped == '%':
                continue
            if pdf_junk_pattern.search(stripped) and not stripped.startswith('/Title'):
                continue
            if stripped.startswith('/Title'):
                continue
            clean_lines.append(stripped)

        # 5. Append extracted titles to clean text if present
        if extracted_titles:
            clean_lines.append("\nADDITIONAL EXTRACTED CONTENT:")
            for t in extracted_titles:
                clean_lines.append(t)

        return "\n".join(clean_lines).strip()

    @classmethod
    def parse_file(cls, filename: str, file_bytes: bytes) -> str:
        """Determines file format, extracts, and sanitizes text content."""
        lower_name = filename.lower()
        extracted = ""
        if lower_name.endswith(".pdf"):
            extracted = cls.extract_text_from_pdf(file_bytes)
            if "%PDF-" in extracted or "/Title" in extracted or "endobj" in extracted:
                extracted = cls.sanitize_pdf_raw_text(extracted)
        elif lower_name.endswith((".png", ".jpg", ".jpeg")):
            extracted = cls.extract_text_from_image(file_bytes)
        elif lower_name.endswith(".docx") or lower_name.endswith(".doc"):
            extracted = cls.extract_text_from_docx(file_bytes)
        elif lower_name.endswith(".txt"):
            extracted = file_bytes.decode("utf-8", errors="ignore")
        else:
            raise ValueError("Unsupported file format. Please upload a PDF, PNG, JPG, DOCX, or TXT file.")

        if "%PDF-" in extracted or "/Title" in extracted or "endobj" in extracted:
            extracted = cls.sanitize_pdf_raw_text(extracted)

        extracted = re.sub(r'[\uf000-\uf8ff]', '', extracted)
        return extracted.strip()

